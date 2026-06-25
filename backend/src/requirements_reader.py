"""Read & analyse requirement documents (PDF / DOCX / Markdown / plain text).

Turns a folder of mixed-format briefs into clean Markdown the diagram agent can
reason over. For every source file it emits a sibling ``<name>.md`` (under an
output dir) and returns the extracted text so a caller can concatenate the whole
corpus into a single requirement description.

Dependency-light by design:
- ``.docx`` is parsed with ``python-docx`` when available, else via a stdlib
  ``zipfile`` fallback (so it works even without the extra dependency).
- ``.pdf`` uses ``pypdf`` when available; otherwise that file is skipped with a
  clear note instead of crashing the whole run.
- ``.md`` / ``.txt`` are read verbatim.
"""

from __future__ import annotations

import html
import logging
import re
import zipfile
from dataclasses import dataclass
from pathlib import Path

logger = logging.getLogger(__name__)

# File extensions we know how to read.
PDF_EXT = {".pdf"}
DOCX_EXT = {".docx"}
TEXT_EXT = {".md", ".markdown", ".txt"}
IMAGE_EXT = {".png", ".jpg", ".jpeg", ".webp", ".gif"}
SUPPORTED_EXT = PDF_EXT | DOCX_EXT | TEXT_EXT | IMAGE_EXT


@dataclass
class ParsedDocument:
    """One parsed source document."""

    source: Path          # original file
    title: str            # human title (file stem)
    text: str             # extracted plain text / markdown
    kind: str             # "pdf" | "docx" | "text" | "image"
    error: str | None = None  # set if extraction failed
    # For image uploads: base64-encoded bytes + MIME type (e.g. "image/png").
    image_b64: str | None = None
    image_mime: str | None = None

    @property
    def ok(self) -> bool:
        if self.kind == "image":
            return self.error is None and bool(self.image_b64)
        return self.error is None and bool(self.text.strip())


# --------------------------------------------------------------------------- #
# Per-format extractors
# --------------------------------------------------------------------------- #
def _extract_docx(path: Path) -> str:
    """Extract text from a .docx, including table cells. Tries python-docx,
    falls back to a stdlib zipfile/XML parse so it works with no extra deps."""
    try:
        import docx  # python-docx

        doc = docx.Document(str(path))
        parts: list[str] = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    except ImportError:
        logger.info("python-docx missing; using zipfile fallback for %s", path.name)
        return _extract_docx_zip(path)


def _extract_docx_zip(path: Path) -> str:
    """Stdlib-only .docx fallback: pull text out of word/document.xml."""
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8", "replace")
    # Paragraph and table-row breaks become newlines; tab markers become spaces.
    xml = xml.replace("</w:p>", "\n").replace("</w:tr>", "\n").replace("<w:tab/>", "\t")
    text = html.unescape(re.sub(r"<[^>]+>", "", xml))
    return "\n".join(ln.rstrip() for ln in text.splitlines() if ln.strip())


def _extract_pdf(path: Path) -> str:
    """Extract text from a PDF with pypdf (lazy import)."""
    try:
        from pypdf import PdfReader
    except ImportError as e:  # pragma: no cover - surfaced as a per-file error
        raise RuntimeError("pypdf not installed (pip install pypdf)") from e

    reader = PdfReader(str(path))
    pages: list[str] = []
    for i, page in enumerate(reader.pages):
        try:
            pages.append(page.extract_text() or "")
        except Exception as e:  # noqa: BLE001 - one bad page shouldn't kill the file
            logger.debug("pdf page %d of %s failed: %s", i, path.name, e)
    return "\n\n".join(p.strip() for p in pages if p.strip())


def _extract_text(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


_IMAGE_MIME: dict[str, str] = {
    ".png": "image/png",
    ".jpg": "image/jpeg",
    ".jpeg": "image/jpeg",
    ".webp": "image/webp",
    ".gif": "image/gif",
}


def _extract_image(path: Path) -> tuple[str, str]:
    """Read an image file and return (base64_string, mime_type)."""
    import base64
    mime = _IMAGE_MIME.get(path.suffix.lower(), "image/png")
    b64 = base64.standard_b64encode(path.read_bytes()).decode("ascii")
    return b64, mime


# --------------------------------------------------------------------------- #
# Public API
# --------------------------------------------------------------------------- #
def parse_file(path: Path) -> ParsedDocument:
    """Parse a single document; never raises — failures land in ``.error``."""
    path = Path(path)
    ext = path.suffix.lower()
    title = path.stem
    try:
        if ext in DOCX_EXT:
            return ParsedDocument(path, title, _extract_docx(path), "docx")
        if ext in PDF_EXT:
            return ParsedDocument(path, title, _extract_pdf(path), "pdf")
        if ext in TEXT_EXT:
            return ParsedDocument(path, title, _extract_text(path), "text")
        if ext in IMAGE_EXT:
            b64, mime = _extract_image(path)
            return ParsedDocument(path, title, "", "image", image_b64=b64, image_mime=mime)
        return ParsedDocument(path, title, "", "unknown",
                              error=f"unsupported extension '{ext}'")
    except Exception as e:  # noqa: BLE001 - report, don't crash the batch
        logger.warning("failed to parse %s: %s", path.name, e)
        return ParsedDocument(path, title, "", ext.lstrip(".") or "unknown", error=str(e))


def _to_markdown(doc: ParsedDocument) -> str:
    """Render a parsed doc as a tidy Markdown file."""
    head = f"# {doc.title}\n\n> Source: `{doc.source.name}` · type: {doc.kind}\n\n"
    if not doc.ok:
        return head + f"**Extraction failed:** {doc.error or 'no text content'}\n"
    return head + "---\n\n" + doc.text.strip() + "\n"


def read_folder(
    folder: str | Path,
    out_dir: str | Path | None = None,
    *,
    write_md: bool = True,
) -> list[ParsedDocument]:
    """Parse every supported file in ``folder``.

    If ``write_md`` is True, writes ``<name>.md`` for each file into ``out_dir``
    (default: ``<folder>/_parsed_md``). Returns the list of ParsedDocument.
    """
    folder = Path(folder)
    if out_dir is None:
        out_dir = folder / "_parsed_md"
    out_dir = Path(out_dir)
    if write_md:
        out_dir.mkdir(parents=True, exist_ok=True)

    files = sorted(
        p for p in folder.iterdir()
        if p.is_file() and p.suffix.lower() in SUPPORTED_EXT
    )
    docs: list[ParsedDocument] = []
    for path in files:
        doc = parse_file(path)
        docs.append(doc)
        status = "ok" if doc.ok else f"FAILED ({doc.error})"
        logger.info("parsed %-45s %s (%d chars)", path.name, status, len(doc.text))
        if write_md:
            md_path = out_dir / f"{path.stem}.md"
            md_path.write_text(_to_markdown(doc), encoding="utf-8")
    return docs


def combine_corpus(docs: list[ParsedDocument], *, max_chars: int = 60_000) -> str:
    """Concatenate successfully-parsed docs into one requirement description,
    capped at ``max_chars`` so it fits comfortably in the model context."""
    chunks: list[str] = []
    for doc in docs:
        if doc.ok:
            chunks.append(f"## {doc.title}\n\n{doc.text.strip()}")
    corpus = "\n\n---\n\n".join(chunks)
    if len(corpus) > max_chars:
        corpus = corpus[:max_chars] + "\n\n[... truncated for length ...]"
    return corpus


if __name__ == "__main__":  # quick CLI: python -m diagram_mcp.requirements_reader <folder>
    import sys

    logging.basicConfig(level=logging.INFO, format="%(message)s")
    target = sys.argv[1] if len(sys.argv) > 1 else "."
    parsed = read_folder(target)
    ok = sum(d.ok for d in parsed)
    print(f"\nParsed {ok}/{len(parsed)} files; .md written to {Path(target) / '_parsed_md'}")
